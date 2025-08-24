import os
import re
import requests
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches
import urllib.parse
import uuid


class MarkdownToWordConverter:
    def __init__(self, input_folder, output_folder):
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.image_counter = 1

    def convert_all_markdown_files(self):
        """主转换函数"""
        if not self.input_folder.exists():
            print(f"错误：输入文件夹不存在 {self.input_folder}")
            return False

        self.output_folder.mkdir(parents=True, exist_ok=True)
        print(f"开始转换: {self.input_folder} → {self.output_folder}")

        success_count = 0
        for root, _, files in os.walk(self.input_folder):
            for file in files:
                if file.lower().endswith('.md'):
                    md_file = Path(root) / file
                    if self._process_markdown_file(md_file):
                        success_count += 1

        print(f"\n转换完成！成功转换 {success_count} 个文件")
        return True

    def _process_markdown_file(self, md_file):
        """处理单个Markdown文件"""
        try:
            # 准备输出目录
            relative_path = md_file.relative_to(self.input_folder)
            output_dir = self.output_folder / relative_path.parent
            output_dir.mkdir(parents=True, exist_ok=True)

            print(f"\n处理文件: {relative_path}")

            # 第一步：下载所有PDF文件
            pdf_urls = self._extract_all_pdf_links(md_file)
            for text, url in pdf_urls:
                self._download_pdf(url, md_file.parent, output_dir)

            # 创建Word文档
            doc = Document()
            with open(md_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()

            # 处理第一行作为一级标题
            if lines:
                first_line = lines[0].strip()
                if first_line:
                    doc.add_heading(first_line, level=1)
                lines = lines[1:]

            current_para = None

            for line in lines:
                line = line.strip()
                if not line:
                    current_para = None
                    continue

                # 第二步：处理下划线
                line = self._process_underline(line)

                # 第三步：处理info块
                if line.startswith(':::'):
                    info_content = self._process_info_block(line)
                    if info_content:
                        p = doc.add_paragraph(info_content)
                        p.style = 'Normal'
                    current_para = None
                    continue

                # 第四步：处理####标题
                if line.startswith('####'):
                    heading_text = line[4:].strip()
                    doc.add_heading(heading_text, level=2)
                    current_para = None
                    continue

                # 第五步：处理**标题**
                if line.startswith('**') and line.endswith('**'):
                    heading_text = line[2:-2].strip()
                    doc.add_heading(heading_text, level=2)
                    current_para = None
                    continue

                # 第六步：处理图片
                image_match = self._extract_image(line)
                if image_match:
                    alt_text, image_path = image_match
                    self._insert_image(doc, alt_text, image_path, md_file.parent)
                    continue

                # 处理普通文本
                if line:
                    if current_para is None:
                        current_para = doc.add_paragraph(line)
                    else:
                        current_para.add_run('\n' + line)

            # 保存Word文档
            doc.save(output_dir / f"{md_file.stem}.docx")
            print("✓ 转换成功")
            return True

        except Exception as e:
            print(f"✗ 转换失败: {str(e)}")
            return False

    def _extract_all_pdf_links(self, md_file):
        """提取文件中所有PDF链接"""
        pdf_urls = []
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        pattern = r'\[(.*?)\]\((.*?\.pdf)\)'
        matches = re.finditer(pattern, content, re.IGNORECASE)
        for match in matches:
            if match.group(2).lower().endswith('.pdf'):
                pdf_urls.append((match.group(1), match.group(2)))

        return pdf_urls

    def _download_pdf(self, url, md_dir, output_dir):
        """下载或复制PDF文件"""
        try:
            if url.startswith(('http://', 'https://')):
                # 下载网络PDF
                response = requests.get(url, timeout=30)
                response.raise_for_status()

                # 从URL获取文件名
                filename = os.path.basename(urllib.parse.urlparse(url).path)
                if not filename.lower().endswith('.pdf'):
                    filename = f"downloaded_{uuid.uuid4().hex[:8]}.pdf"

                output_path = output_dir / filename
                with open(output_path, 'wb') as f:
                    f.write(response.content)

                print(f"  已下载PDF: {filename}")
            else:
                # 复制本地PDF
                src_path = md_dir / url
                if src_path.exists():
                    dst_path = output_dir / src_path.name
                    shutil.copy2(src_path, dst_path)
                    print(f"  已复制PDF: {src_path.name}")
                else:
                    print(f"  警告: PDF文件不存在: {src_path}")
        except Exception as e:
            print(f"  处理PDF出错: {str(e)}")

    def _process_underline(self, line):
        """处理下划线文本"""

        def replace_underline(match):
            return match.group(1)

        return re.sub(r'<u>(.+?)</u>', replace_underline, line)

    def _process_info_block(self, line):
        """处理info块"""
        if line.startswith(':::info') and line.endswith(':::'):
            return line[7:-3].strip()
        elif line.startswith(':::'):
            return line[3:].strip()
        return line

    def _extract_image(self, line):
        """提取图片链接"""
        pattern = r'!\[(.*?)\]\((.*?)\)'
        match = re.search(pattern, line)
        if match:
            return match.groups()
        return None

    def _insert_image(self, doc, alt_text, image_path, md_dir):
        """插入图片到Word"""
        if image_path.startswith(('http://', 'https://')):
            print(f"  跳过网络图片: {image_path}")
            return

        full_image_path = md_dir / image_path

        if full_image_path.exists():
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = 1  # 居中
                run = paragraph.add_run()

                run.add_picture(str(full_image_path), width=Inches(5))

                if alt_text and alt_text.strip():
                    caption = doc.add_paragraph(f"图{self.image_counter}: {alt_text}")
                    caption.alignment = 1  # 居中
                    caption.runs[0].italic = True
                    self.image_counter += 1

                print(f"  已插入图片: {image_path}")

            except Exception as e:
                print(f"  插入图片失败: {str(e)}")
        else:
            print(f"  警告: 图片文件不存在: {full_image_path}")


if __name__ == "__main__":
    # 配置路径
    input_dir = r"docs"  # 替换为您的输入目录
    output_dir = r"output"  # 替换为您的输出目录

    # 执行转换
    converter = MarkdownToWordConverter(input_dir, output_dir)
    if converter.convert_all_markdown_files():
        print("\n处理完成！")
    else:
        print("\n处理过程中出现错误")
